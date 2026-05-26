"""НейроБокс — Модуль генерации документов: PDF, DOCX, XLSX, TXT, CSV, Markdown."""
import io
import re
import time

import structlog
from aiogram import F, Router, types
from aiogram.filters import BaseFilter, Command
from aiogram.utils.keyboard import InlineKeyboardBuilder

from shared.db.database import get_pool
from shared.domain.credits import (
    get_or_create_user,
    refund_spend_credits,
    spend_credits,
)
from shared.domain.telemetry import elapsed_ms, log_ai_success
from shared.providers.openai_text import generate_text

router = Router()
_log = structlog.get_logger()

# Стоимость в кредитах
DOCGEN_CR = {
    "txt": 6, "md": 6, "csv": 6,
    "docx": 10, "pdf": 12, "xlsx": 10,
}
DOCGEN_MODEL = "gpt-5-nano"

_SYSTEM_PROMPT = (
    "Ты — профессиональный копирайтер-редактор. "
    "Генерируй структурированный текст с заголовками (## Заголовок), подзаголовками (### Подзаголовок), "
    "нумерованными и ненумерованными списками. Стиль: деловой, без воды, конкретный. "
    "Язык: русский (если пользователь не указал иной). "
    "Объём: адекватный типу документа (резюме 1 стр, бизнес-план 5-8 стр). "
    "НЕ пиши вступление типа 'Конечно, вот документ'. Начинай сразу с содержимого."
)


def _retry_kb(user_id: int, fmt: str):
    b = InlineKeyboardBuilder()
    b.row(types.InlineKeyboardButton(text="🔄 Попробовать снова", callback_data=f"docgen_retry:{user_id}:{fmt}"))
    b.row(types.InlineKeyboardButton(text="◀️ Меню", callback_data="main_menu"))
    return b.as_markup()


def _after_kb(user_id: int):
    b = InlineKeyboardBuilder()
    b.row(
        types.InlineKeyboardButton(text="📄 Ещё документ", callback_data="screen_docgen"),
        types.InlineKeyboardButton(text="🗣 Озвучить", callback_data=f"docgen_tts:{user_id}"),
    )
    b.row(types.InlineKeyboardButton(text="◀️ Меню", callback_data="main_menu"))
    return b.as_markup()


def _transliterate(text: str) -> str:
    """Транслитерация кириллицы для имени файла."""
    mapping = {
        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo',
        'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
        'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
        'ф': 'f', 'х': 'kh', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'sch',
        'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
    }
    result = []
    for ch in text.lower():
        if ch in mapping:
            result.append(mapping[ch])
        elif ch.isalnum() or ch in '-_':
            result.append(ch)
        elif ch == ' ':
            result.append('_')
    name = ''.join(result)[:60].strip('_')
    return name or 'document'


def _make_filename(task: str, fmt: str) -> str:
    """Осмысленное имя файла из описания задачи."""
    words = task.split()[:5]
    short = ' '.join(words)
    return f"{_transliterate(short)}.{fmt}"


def _build_txt(content: str) -> bytes:
    return content.encode('utf-8')


def _build_md(content: str) -> bytes:
    return content.encode('utf-8')


def _build_csv(content: str) -> bytes:
    """Пытаемся структурировать как CSV, иначе — как есть."""
    return content.encode('utf-8')


def _build_docx(content: str, title: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    if title:
        doc.add_heading(title, level=0)
    for line in content.split('\n'):
        line = line.rstrip()
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+[\.\)] ', line):
            doc.add_paragraph(re.sub(r'^\d+[\.\)] ', '', line), style='List Number')
        elif line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(content: str, title: str = "") -> bytes:
    """Генерация PDF через reportlab с поддержкой кириллицы."""
    try:
        import os

        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        # Попытка зарегистрировать кириллический шрифт
        font_name = 'DejaVuSans'
        font_paths = [
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/dejavu/DejaVuSans.ttf',
            '/app/fonts/DejaVuSans.ttf',
        ]
        font_registered = False
        for fp in font_paths:
            if os.path.exists(fp):
                pdfmetrics.registerFont(TTFont(font_name, fp))
                font_registered = True
                break
        if not font_registered:
            font_name = 'Helvetica'

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=14)
        h1_style = ParagraphStyle('H1', parent=styles['Heading1'], fontName=font_name, fontSize=18, leading=22)
        h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontName=font_name, fontSize=14, leading=18)
        h3_style = ParagraphStyle('H3', parent=styles['Heading3'], fontName=font_name, fontSize=12, leading=16)

        story = []
        if title:
            story.append(Paragraph(title, h1_style))
            story.append(Spacer(1, 6*mm))
        for line in content.split('\n'):
            line = line.rstrip()
            if line.startswith('### '):
                story.append(Paragraph(line[4:], h3_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], h2_style))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], h1_style))
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f"• {line[2:]}", body_style))
            elif line.strip():
                story.append(Paragraph(line, body_style))
            else:
                story.append(Spacer(1, 3*mm))
        doc.build(story)
        return buf.getvalue()
    except ImportError:
        # Fallback: отправить как TXT если reportlab недоступен
        return content.encode('utf-8')


def _build_xlsx(content: str) -> bytes:
    """Пытаемся распарсить табличные данные в XLSX."""
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Документ"
        for i, line in enumerate(content.split('\n'), 1):
            line = line.strip()
            if not line:
                continue
            # Попробовать разделить по | (markdown-таблица) или по ;/,
            if '|' in line:
                cells = [c.strip() for c in line.split('|') if c.strip() and c.strip() != '---']
            elif ';' in line:
                cells = [c.strip() for c in line.split(';')]
            elif ',' in line and line.count(',') >= 2:
                cells = [c.strip() for c in line.split(',')]
            else:
                cells = [line]
            for j, cell in enumerate(cells, 1):
                ws.cell(row=i, column=j, value=cell)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except ImportError:
        return content.encode('utf-8')


BUILDERS = {
    'txt': _build_txt, 'md': _build_md, 'csv': _build_csv,
    'docx': _build_docx, 'pdf': _build_pdf, 'xlsx': _build_xlsx,
}


async def _log_docgen(user_id: int, doc_type: str, fmt: str, title: str, tokens: int, cr: int, size: int):
    try:
        pool = await get_pool()
        async with pool.acquire() as conn:
            await conn.execute(
                "INSERT INTO document_generations (user_id, doc_type, format, title, tokens_used, credits_charged, file_size_bytes) "
                "VALUES ($1, $2, $3, $4, $5, $6, $7)",
                user_id, doc_type, fmt, (title or "")[:300], tokens, cr, size)
    except Exception:
        pass


async def _generate_document(user_id: int, task: str, fmt: str, message: types.Message):
    """Основная логика генерации документа."""
    cr = DOCGEN_CR.get(fmt, 10)
    spend = await spend_credits(user_id, DOCGEN_MODEL, f"docgen_{fmt}: {task[:40]}", cost_override=cr)
    if not spend["ok"]:
        if spend.get("message"):
            await message.answer(spend["message"])
            return
        from services.bot.utils.paywall import smart_paywall_message
        text, kb = await smart_paywall_message(f"Генерация документа ({fmt.upper()})", cr, user_id)
        await message.answer(text, reply_markup=kb)
        return

    await message.chat.do("upload_document")
    started = time.monotonic()

    # Генерация контента через LLM
    fmt_hint = {
        'xlsx': "Ответь в формате таблицы с разделителем |. Первая строка — заголовки.",
        'csv': "Ответь в формате CSV с разделителем ;. Первая строка — заголовки.",
    }.get(fmt, "")
    prompt = f"Создай документ: {task}\n\n{fmt_hint}".strip()
    result = await generate_text(prompt, DOCGEN_MODEL, history=None, system_prompt=_SYSTEM_PROMPT)

    if not result.get("ok"):
        await refund_spend_credits(user_id, spend, "ошибка генерации документа")
        await message.answer(
            f"❌ {result.get('error', 'Ошибка генерации')}",
            reply_markup=_retry_kb(user_id, fmt))
        return

    content = (result.get("text") or "").strip()
    if not content:
        await refund_spend_credits(user_id, spend, "пустой документ")
        await message.answer("❌ Не удалось сгенерировать контент.", reply_markup=_retry_kb(user_id, fmt))
        return

    # Построение файла
    builder = BUILDERS.get(fmt, _build_txt)
    title = task[:100]
    try:
        if fmt in ('docx', 'pdf'):
            file_bytes = builder(content, title)
        else:
            file_bytes = builder(content)
    except Exception as e:
        _log.error("docgen_build_error", error=str(e)[:200], fmt=fmt)
        await refund_spend_credits(user_id, spend, "ошибка сборки файла")
        await message.answer("❌ Ошибка сборки файла.", reply_markup=_retry_kb(user_id, fmt))
        return

    if len(file_bytes) > 20 * 1024 * 1024:
        await refund_spend_credits(user_id, spend, "файл слишком большой")
        await message.answer("❌ Документ слишком большой (>20 МБ). Сократи описание.")
        return

    filename = _make_filename(task, fmt)
    remaining = spend.get("bought", 0) + spend.get("free", 0)
    tokens = result.get("tokens_out", 0) or len(content) // 4

    await log_ai_success(user_id, "docgen", DOCGEN_MODEL, task[:200], cr, elapsed_ms(started))
    await _log_docgen(user_id, "from_scratch", fmt, title, tokens, cr, len(file_bytes))

    # Сохранить текст в Redis для TTS
    try:
        from shared.redis.store import _get_redis
        r = await _get_redis()
        if r:
            await r.set(f"docgen_text:{user_id}", content[:4000], ex=600)
    except Exception:
        pass

    from aiogram.types import BufferedInputFile
    doc_file = BufferedInputFile(file_bytes, filename=filename)
    await message.answer_document(
        document=doc_file,
        caption=f"📄 <b>{title[:80]}</b>\n\n<i>−{cr} CR | Остаток: {remaining} CR</i>",
        reply_markup=_after_kb(user_id))


# ── Экраны и обработчики ──

@router.callback_query(F.data == "screen_docgen")
async def cb_screen_docgen(cb: types.CallbackQuery):
    await cb.answer()
    b = InlineKeyboardBuilder()
    for fmt, label in [("pdf", "📕 PDF"), ("docx", "📘 DOCX"), ("xlsx", "📗 XLSX"),
                       ("txt", "📝 TXT"), ("csv", "📊 CSV"), ("md", "📓 Markdown")]:
        cr = DOCGEN_CR.get(fmt, 6)
        b.button(text=f"{label} ({cr} CR)", callback_data=f"docgen_fmt:{fmt}")
    b.adjust(2)
    b.row(types.InlineKeyboardButton(text="◀️ Назад", callback_data="more_menu"))
    await cb.message.answer(
        "📄 <b>Генерация документов</b>\n\n"
        "Выбери формат, затем опиши что нужно создать:\n"
        "резюме, договор, бизнес-план, таблицу...",
        reply_markup=b.as_markup())


@router.callback_query(F.data.startswith("docgen_fmt:"))
async def cb_docgen_fmt(cb: types.CallbackQuery):
    await cb.answer()
    fmt = cb.data.replace("docgen_fmt:", "")
    cr = DOCGEN_CR.get(fmt, 6)
    # Сохраняем выбранный формат в Redis
    try:
        from shared.redis.store import _get_redis
        r = await _get_redis()
        if r:
            await r.set(f"docgen_fmt:{cb.from_user.id}", fmt, ex=300)
    except Exception:
        pass
    b = InlineKeyboardBuilder()
    b.row(types.InlineKeyboardButton(text="◀️ Назад", callback_data="screen_docgen"))
    await cb.message.answer(
        f"📄 Формат: <b>{fmt.upper()}</b> ({cr} CR)\n\n"
        "Опиши документ. Примеры:\n"
        "• <code>Резюме frontend-разработчика Иван Петров, 5 лет опыта</code>\n"
        "• <code>Бизнес-план кофейни на 20 посадочных мест</code>\n"
        "• <code>Таблица расходов на маркетинг по месяцам</code>",
        reply_markup=b.as_markup())


@router.callback_query(F.data.startswith("docgen_retry:"))
async def cb_docgen_retry(cb: types.CallbackQuery):
    parts = cb.data.split(":")
    if len(parts) < 3:
        await cb.answer("Отправь описание заново.")
        return
    await cb.answer("🔄 Повторяю...")
    await cb.message.answer("Отправь описание документа ещё раз.")


@router.callback_query(F.data.startswith("docgen_tts:"))
async def cb_docgen_tts(cb: types.CallbackQuery):
    """Озвучить последний сгенерированный документ."""
    await cb.answer()
    user_id = cb.from_user.id
    try:
        from shared.redis.store import _get_redis
        r = await _get_redis()
        if r:
            text = await r.get(f"docgen_text:{user_id}")
            if text:
                text = text.decode() if isinstance(text, bytes) else text
                await r.set(f"tts_pending:{user_id}", text[:4000], ex=300)
                await cb.message.answer(
                    "🗣 Текст документа готов к озвучке.\n"
                    "Выбери голос в разделе 🎵 Аудио → 🗣 Озвучить текст.")
                return
    except Exception:
        pass
    await cb.message.answer("Текст документа не найден. Сгенерируй документ заново.")


@router.message(Command("gendoc"))
async def cmd_gendoc(message: types.Message):
    """Команда /gendoc формат описание — генерация документа."""
    user_id = message.from_user.id
    await get_or_create_user(user_id, message.from_user.username, message.from_user.first_name)
    args = (message.text or "").split(maxsplit=2)
    if len(args) < 3:
        await message.answer(
            "📄 Формат: <code>/gendoc pdf Резюме разработчика</code>\n"
            "Форматы: pdf, docx, xlsx, txt, csv, md")
        return
    fmt = args[1].lower().strip('.')
    if fmt not in BUILDERS:
        await message.answer(f"❌ Неизвестный формат <b>{fmt}</b>. Доступные: pdf, docx, xlsx, txt, csv, md")
        return
    task = args[2].strip()
    await _generate_document(user_id, task, fmt, message)


# ── Async Filter для перехвата текста ──

class HasDocgenFormat(BaseFilter):
    """Фильтр: True только если у пользователя выбран формат документа в Redis."""
    async def __call__(self, message: types.Message) -> bool | dict:
        try:
            from shared.redis.store import _get_redis
            r = await _get_redis()
            if not r:
                return False
            fmt = await r.get(f"docgen_fmt:{message.from_user.id}")
            if not fmt:
                return False
            fmt = fmt.decode() if isinstance(fmt, bytes) else fmt
            await r.delete(f"docgen_fmt:{message.from_user.id}")
            return {"docgen_fmt": fmt}
        except Exception:
            return False


@router.message(HasDocgenFormat(), F.text & ~F.text.startswith("/"))
async def handle_docgen_text(message: types.Message, docgen_fmt: str = "txt"):
    """Перехват текста после выбора формата документа."""
    user_id = message.from_user.id
    task = (message.text or "").strip()
    if not task or len(task) < 5:
        await message.answer("Опиши документ подробнее (минимум 5 символов).")
        return
    await get_or_create_user(user_id, message.from_user.username, message.from_user.first_name)
    await _generate_document(user_id, task, docgen_fmt, message)
